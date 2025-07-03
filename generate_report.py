from docx import Document
from resume_parser import extract_text_from_docx, extract_text_from_pdf, parse_resume
from job_description_parser import parse_job_description
from resume_scorer import score_resume, generate_feedback
import os

# === CONFIGURATION ===
cv_filename = "Talha_Sarfraz_ATS_Resume.docx"  # Change to your file name
job_description = "full stack developer at XYZ Corp."
job_responsibilities = "Develop and maintain web applications."
job_experience = "2+ years in web development."
job_skills = "html , css , javascrpt , express.js , nodejs , sql databases"
job_education = "Bachelor's in Computer Science or related fieds"

# === LOAD CV ===
suffix = os.path.splitext(cv_filename)[1].lower()
if suffix == ".pdf":
    text = extract_text_from_pdf(cv_filename)
elif suffix == ".docx":
    text = extract_text_from_docx(cv_filename)
else:
    raise ValueError("Unsupported file type")

# === PARSE RESUME ===
resume_data = parse_resume(text)

# === PARSE JOB DESCRIPTION ===
job_desc = parse_job_description(
    job_description,
    job_responsibilities,
    job_experience,
    job_skills,
    job_education
)

# === SCORE RESUME ===
score = score_resume(resume_data, job_desc)

# === GENERATE FEEDBACK ===
feedback = generate_feedback(resume_data, job_desc)

# === EXTRACT CANDIDATE NAME ===
entities = resume_data.get('entities', [])
candidate_name = entities[0] if entities else 'Unknown Candidate'

# === CREATE DOCX REPORT ===
doc = Document()
doc.add_heading('Resume Analysis Report', 0)


doc.add_paragraph(candidate_name)

# Add Resume Score
doc.add_heading('Resume Score', level=1)
for k, v in score.items():
    if k != 'sections':
        doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
    else:
        doc.add_paragraph("Sections Detected:")
        for section, present in v.items():
            doc.add_paragraph(f"  - {section}: {'Yes' if present else 'No'}", style='List Bullet')

# Add Full CV Text Section
doc.add_heading('CV Text', level=1)
doc.add_paragraph(resume_data.get('full_text', ''))

# Remove Parsed Resume Data section

# Add Feedback
doc.add_heading('Feedback', level=1)
for line in feedback.split('. '):
    if line.strip():
        doc.add_paragraph(line.strip(), style='List Bullet')

# Save the document
output_filename = "resume_analysis_report.docx"
doc.save(output_filename)
print(f"Report saved as {output_filename}")